from docx import Document

def extract_text_from_word(file_path):
    # document = Document(file_path)
    # text = ""
    # for paragraph in document.paragraphs:
    #     text += paragraph.text + "\n"
    # return text.strip()
    document = Document(file_path)
    sections = []
    current_title = None
    current_content = []
    
    for paragraph in document.paragraphs:
        text = paragraph.text.strip()
        if text.isupper():  # Assuming titles are in uppercase
            if current_title:
                sections.append({"title": current_title, "content": "\n".join(current_content)})
            current_title = text
            current_content = []
        else:
            current_content.append(text)
    
    if current_title:
        sections.append({"title": current_title, "content": "\n".join(current_content)})
    
    return sections

# Example usage
word_file_path = ".\AltoK10_Owner's_Manual.docx"
text_content = extract_text_from_word(word_file_path)
print(text_content)
